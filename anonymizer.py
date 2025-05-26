import os
import re
import json
import random
import base64
import logging
import zipfile

from faker import Faker
import spacy
from cryptography.fernet import Fernet

try:
    import pandas as pd
    import openpyxl
except ImportError:
    pd = None

try:
    from text_deidentification.deidentifier import DeidentifiedText
except ImportError:
    DeidentifiedText = None

from docx import Document
import fitz  # PyMuPDF

logging.basicConfig(
    format="%(asctime)s %(levelname)s: %(message)s", level=logging.INFO
)
fake = Faker("ru_RU")

TXT_BEGIN = b"\n--MAPPING-START--\n"
TXT_END   = b"\n--MAPPING-END--\n"

# load spaCy models
try:
    nlp_ru = spacy.load("ru_core_news_sm")
except Exception:
    nlp_ru = None
    logging.warning("spaCy model 'ru_core_news_sm' not found — русский NER отключен.")
try:
    nlp_en = spacy.load("en_core_web_sm")
except Exception:
    nlp_en = None
    logging.warning("spaCy model 'en_core_web_sm' not found — английский NER отключен.")


def load_patterns():
    return {
        "fio": re.compile(r"\b(?:[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?\s+){1,2}[А-ЯЁ][а-яё]+\b"),
        "email": re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"),
        "phone": re.compile(r"(?:\+7|8)\s*\d{3}[\s-]?\d{3}[\s-]?\d{2}[\s-]?\d{2}"),
        "passport": re.compile(r"\b\d{4}\s?\d{6}\b"),
        "snils": re.compile(r"\b\d{3}-\d{3}-\d{3}\s?\d{2}\b"),
        "inn": re.compile(r"\b\d{10}\b|\b\d{12}\b"),
        "card": re.compile(r"\b(?:\d{4}[\s-]?){3}\d{4}\b"),
        "address": re.compile(
            r"\b(?:ул(?:\.|ица)?|проспект|пр-т|пл(?:\.|ощадь)?|бульвар|пер(?:\.|евоз)?)"
            r"\s+[А-ЯЁ][\w\-\s]+?\s*\d+\w?(?:,\s*\d+)?\b"
        ),
    }


def mask_token(tok: str) -> str:
    if len(tok) <= 2:
        return "*" * len(tok)
    return tok[0] + "*" * (len(tok) - 2) + tok[-1]


def encrypt_token(tok: str, key: str) -> str:
    return Fernet(key.encode()).encrypt(tok.encode()).decode()


def chaos_keystream(length: int, seed: float, r: float) -> bytes:
    x = seed
    ks = bytearray(length)
    for i in range(length):
        x = r * x * (1 - x)
        ks[i] = int(x * 256) & 0xFF
    return bytes(ks)


def chaos_encrypt(data: bytes, seed: float, r: float) -> str:
    ks = chaos_keystream(len(data), seed, r)
    ct = bytes(b ^ k for b, k in zip(data, ks))
    return base64.urlsafe_b64encode(ct).decode()


def chaos_decrypt(token: str, seed: float, r: float) -> bytes:
    ct = base64.urlsafe_b64decode(token)
    ks = chaos_keystream(len(ct), seed, r)
    return bytes(c ^ k for c, k in zip(ct, ks))


def initialize_parameters(method: str, params: dict):
    if method == "encrypt":
        params["key"] = Fernet.generate_key().decode()
    elif method == "chaos":
        params["seed"], params["r"] = random.random(), 3.99
    elif method == "pseudonymize":
        params["seed"] = random.randint(0, 2**31 - 1)


def anonymize_regex(text: str, method: str, patterns: dict, cache: dict, mapping: dict) -> str:
    for cat, pat in patterns.items():
        def repl(m):
            orig = m.group(0)
            if orig in cache:
                return cache[orig]
            if method == "pseudonymize":
                if cat == "fio":
                    anon = fake.name()
                elif cat == "email":
                    anon = fake.email()
                elif cat == "phone":
                    anon = fake.phone_number()
                elif cat == "address":
                    anon = fake.address()
                else:
                    anon = f"TOKEN_{len(cache)+1}"
            elif method == "mask":
                anon = mask_token(orig)
            elif method == "encrypt":
                anon = encrypt_token(orig, mapping["parameters"]["key"])
            elif method == "chaos":
                anon = chaos_encrypt(orig.encode("utf-8"),
                                     mapping["parameters"]["seed"],
                                     mapping["parameters"]["r"])
            else:
                anon = orig
            cache[orig] = anon
            mapping["mappings"].append({"original": orig, "anonymized": anon})
            return anon
        text = pat.sub(repl, text)
    return text


def anonymize_spacy_entities(text: str, cache: dict, mapping: dict, nlp_model, method: str) -> str:
    if not nlp_model:
        logging.error("NER-модель не загружена — пропуск spaCy-анонимизации")
        return text
    doc = nlp_model(text)
    for ent in reversed(doc.ents):
        if ent.label_ in ("PER", "PERSON"):
            orig = ent.text
            if orig in cache:
                anon = cache[orig]
            else:
                if method == "pseudonymize":
                    anon = fake.name()
                elif method == "mask":
                    anon = mask_token(orig)
                elif method == "encrypt":
                    anon = encrypt_token(orig, mapping["parameters"]["key"])
                elif method == "chaos":
                    anon = chaos_encrypt(orig.encode("utf-8"),
                                         mapping["parameters"]["seed"],
                                         mapping["parameters"]["r"])
                else:
                    anon = orig
                cache[orig] = anon
                mapping["mappings"].append({"original": orig, "anonymized": anon})
            text = text[:ent.start_char] + anon + text[ent.end_char:]
    return text


def _embed_mapping_txt(path: str, mapping: dict):
    b64 = base64.b64encode(json.dumps(mapping, ensure_ascii=False).encode())
    with open(path, "ab") as f:
        f.write(TXT_BEGIN); f.write(b64); f.write(TXT_END)


def _extract_mapping_txt(raw: bytes):
    if TXT_BEGIN not in raw or TXT_END not in raw:
        raise ValueError("Mapping not found")
    pre, rest = raw.split(TXT_BEGIN, 1)
    b64, _ = rest.split(TXT_END, 1)
    return pre.decode("utf-8"), json.loads(base64.b64decode(b64))


def anonymize_file(path: str,
                   method: str,
                   categories: list = None,
                   lang: str = "ru",
                   chaos_seed: float = None,
                   chaos_r: float = None,
                   pseudo_seed: int = None):
    base, ext = os.path.splitext(path)
    out = f"{base}_anon{ext}"
    mapping = {"method": method, "parameters": {}, "mappings": []}
    initialize_parameters(method, mapping["parameters"])

    if method == "chaos" and chaos_seed is not None:
        mapping["parameters"]["seed"], mapping["parameters"]["r"] = chaos_seed, chaos_r
    if method == "pseudonymize":
        seed = (mapping["parameters"]["seed"] if pseudo_seed is None else pseudo_seed)
        mapping["parameters"]["seed"] = seed
        fake.seed_instance(seed)

    patterns = load_patterns()
    if categories:
        patterns = {c: patterns[c] for c in categories if c in patterns}

    nlp_model = nlp_ru if lang == "ru" else nlp_en
    cache = {}

    # TXT/CSV
    if ext.lower() in (".txt", ".csv"):
        raw = open(path, "rb").read()
        text = raw.decode("utf-8")
        if method == "ner" and DeidentifiedText:
            res = DeidentifiedText().run(text)
        elif method == "ner":
            res = anonymize_spacy_entities(text, cache, mapping, nlp_model, method)
        else:
            res = anonymize_regex(text, method, patterns, cache, mapping)
        with open(out, "wb") as f:
            f.write(res.encode("utf-8"))
        _embed_mapping_txt(out, mapping)
        logging.info(f"TXT → {out}")
        return

    # XLSX
    if ext.lower() == ".xlsx":
        if pd is None:
            logging.error(".xlsx требует pandas"); return
        df = pd.read_excel(path, header=None, dtype=str)
        def anon_cell(c):
            s = str(c)
            if method == "ner":
                return anonymize_spacy_entities(s, cache, mapping, nlp_model, method)
            return anonymize_regex(s, method, patterns, cache, mapping)
        df_out = df.applymap(anon_cell)
        df_out.to_excel(out, index=False, header=False)
        wb = openpyxl.load_workbook(out)
        ws = wb.create_sheet("_mapping")
        ws["A1"] = json.dumps(mapping, ensure_ascii=False)
        ws.sheet_state = "hidden"
        wb.save(out)
        logging.info(f"XLSX → {out}")
        return

    # DOCX
    if ext.lower() == ".docx":
        doc = Document(path)
        for p in doc.paragraphs:
            for run in p.runs:
                txt = run.text
                new = anonymize_spacy_entities(txt, cache, mapping, nlp_model, method) if method == "ner" \
                      else anonymize_regex(txt, method, patterns, cache, mapping)
                run.text = new
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    txt = cell.text
                    new = anonymize_spacy_entities(txt, cache, mapping, nlp_model, method) if method == "ner" \
                          else anonymize_regex(txt, method, patterns, cache, mapping)
                    cell.text = new
        doc.save(out)
        with zipfile.ZipFile(out, "a") as zf:
            zf.writestr("custom_mapping.json", json.dumps(mapping, ensure_ascii=False))
        logging.info(f"DOCX → {out}")
        return

    # PDF
    if ext.lower() == ".pdf":
        doc = fitz.open(path)
        for page in doc:
            text = page.get_text("text")
            new = anonymize_spacy_entities(text, cache, mapping, nlp_model, method) if method == "ner" \
                  else anonymize_regex(text, method, patterns, cache, mapping)
            page.add_redact_annot(page.rect, fill=(1,1,1))
            page.apply_redactions()
            page.insert_textbox(page.rect, new, fontsize=11)
        meta = doc.metadata
        meta["subject"] = base64.b64encode(json.dumps(mapping).encode()).decode()
        doc.set_metadata(meta)
        doc.save(out, garbage=4, deflate=True)
        logging.info(f"PDF → {out}")
        return

    logging.error(f"Unsupported extension: {ext}")


def deanonymize_file(path: str):
    base, ext = os.path.splitext(path)
    out = f"{base}_restored{ext}"

    # TXT/CSV
    if ext.lower() in (".txt", ".csv"):
        raw = open(path, "rb").read()
        text, mapping = _extract_mapping_txt(raw)
        if mapping["method"] == "chaos":
            s, r = mapping["parameters"]["seed"], mapping["parameters"]["r"]
            for m in mapping["mappings"]:
                m["original"] = chaos_decrypt(m["anonymized"], s, r).decode("utf-8")
        rev = {m["anonymized"]: m["original"] for m in mapping["mappings"]}
        for a, o in rev.items():
            text = text.replace(a, o)
        with open(out, "w", encoding="utf-8") as f:
            f.write(text)
        logging.info(f"Restored TXT → {out}")
        return

    # XLSX
    if ext.lower() == ".xlsx" and pd:
        wb = openpyxl.load_workbook(path, data_only=True)
        if "_mapping" not in wb.sheetnames:
            logging.error("Mapping sheet missing"); return
        mapping = json.loads(wb["_mapping"]["A1"].value)
        wb.remove(wb["_mapping"]); wb.save(out)
        df = pd.read_excel(out, header=None, dtype=str)
        if mapping["method"] == "chaos":
            s, r = mapping["parameters"]["seed"], mapping["parameters"]["r"]
            for m in mapping["mappings"]:
                m["original"] = chaos_decrypt(m["anonymized"], s, r).decode("utf-8")
        rev = {m["anonymized"]: m["original"] for m in mapping["mappings"]}
        df_out = df.applymap(lambda c: rev.get(str(c), str(c)))
        df_out.to_excel(out, index=False, header=False)
        logging.info(f"Restored XLSX → {out}")
        return

    # DOCX
    if ext.lower() == ".docx":
        with zipfile.ZipFile(path, "r") as zf:
            mapping = json.loads(zf.read("custom_mapping.json").decode())
        doc = Document(path)
        if mapping["method"] == "chaos":
            s, r = mapping["parameters"]["seed"], mapping["parameters"]["r"]
            for m in mapping["mappings"]:
                m["original"] = chaos_decrypt(m["anonymized"], s, r).decode("utf-8")
        rev = {m["anonymized"]: m["original"] for m in mapping["mappings"]}
        for p in doc.paragraphs:
            for run in p.runs:
                txt = run.text
                for a, o in rev.items():
                    txt = txt.replace(a, o)
                run.text = txt
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    txt = cell.text
                    for a, o in rev.items():
                        txt = txt.replace(a, o)
                    cell.text = txt
        doc.save(out)
        logging.info(f"Restored DOCX → {out}")
        return

    # PDF
    if ext.lower() == ".pdf":
        doc = fitz.open(path)
        subj = doc.metadata.get("subject", "")
        try:
            mapping = json.loads(base64.b64decode(subj).decode())
        except Exception:
            logging.error("No mapping metadata"); return
        if mapping["method"] == "chaos":
            s, r = mapping["parameters"]["seed"], mapping["parameters"]["r"]
            for m in mapping["mappings"]:
                m["original"] = chaos_decrypt(m["anonymized"], s, r).decode("utf-8")
        rev = {m["anonymized"]: m["original"] for m in mapping["mappings"]}
        for page in doc:
            text = page.get_text("text")
            for a, o in rev.items():
                text = text.replace(a, o)
            page.clean_contents()
            page.insert_text((0, 0), text)
        doc.save(out)
        logging.info(f"Restored PDF → {out}")
        return

    logging.error(f"Unsupported extension: {ext}")
